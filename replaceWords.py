#-*- coding:utf-8 -*-

from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

#获取docx中表格数据
def get_data_docx(path,table_num):

    doc = Document(path)
    style = doc.styles['Normal']
    font = style.font
    font.name = u'宋体'
    font.size = Pt(10.5)

    t = doc.tables[table_num]

    cellv = t.cell(0, 6).text
    # t = doc.tables[table_num]

    # row = table.add_row().cells
    # p = row[0].add_paragraph('left justified text')
    # p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # p = row[1].add_paragraph('right justified text')
    # p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    print cellv
    if '7#' in cellv:
        cellv = cellv.replace('7#','6#')
    print cellv
    t.cell(0,6).text = cellv
    # t.cell(0,6).alignment = WD_TABLE_ALIGNMENT.RIGHT



    # p = t.cell(0,6).add_paragraph(cellv)
    # p.alignment = WD_TABLE_ALIGNMENT.CENTER
    doc.save(path)

    return cellv
#将计算结果写入docx文档中
# def write_to_docx(path,textshow):
#     doc = Document(path)
#     style = doc.styles['Normal']
#     font = style.font
#     font.name = u'新宋体'
#     font.size = Pt(12)
# #
#     for tablenum in range(len(doc.tables)):
#         t = doc.tables[tablenum]
#         get_value = get_data_docx(path, tablenum)
#         t.cell(r,13).text = str("6#楼1单元301")
#
#     doc.save(path)



path = 'source.docx'
get_data_docx(path,0)
