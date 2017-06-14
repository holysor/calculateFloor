#coding:utf-8

from docx import Document
from docx.shared import Pt
import numpy as np
import time
import logging
import os
from Tkinter import *
import sys
import threading


start_time = time.strftime('%Y-%m-%d-%H-%M', time.localtime(time.time()))
pathlog = os.getcwd()+os.sep+'Log'+os.sep+start_time
if not os.path.exists(pathlog):
    os.makedirs(pathlog)

#相关信息存于日志中
logging.basicConfig(level=logging.DEBUG,
                    format='%(asctime)s %(filename)s[line:%(lineno)d] %(levelname)s %(message)s',
                    datefmt='%a, %d %b %Y %H:%M:%S',
                    filename=pathlog+os.sep+'calculatefromWord.log',
                    filemode='w')

console = logging.StreamHandler()
console.setLevel(logging.INFO)
formatter = logging.Formatter('%(name)-12s: %(levelname)-8s %(message)s')
console.setFormatter(formatter)
logging.getLogger('').addHandler(console)

#将写word操作,置于线程中执行,不阻塞GUI
class ThreadDoc(threading.Thread):
    def __init__(self,target,args,text,button):
        threading.Thread.__init__(self)
        self.target = target
        self.args = args
        self.text = text
        self.button = button
    def run(self):
        logging.info('开始时间:' + start_time)
        t1 = time.time()
        try:
            self.target(self.args,self.text)
        except:
            info = sys.exc_info()
            self.text.insert(END, str(info[0]) + ':' + str(info[1]) + '\n')
            self.text.insert(END, '无法正常运行(请先查看表格是否关闭)！\n')
            logging.info(str(info[0]) + ':' + str(info[1]))

        t2 = time.time()
        now = time.strftime('%Y-%m-%d-%H_%M_%S', time.localtime(time.time()))
        self.text.insert(END, '\n结束时间:' + now + ', 用时:' + str(int(t2 - t1)) + 's\n')
        self.text.insert(END,"=======================================================")
        logging.info('结束时间:' + now + ',用时:' + str(int(t2 - t1)) + 's')
        self.button['text'] = '开始'

#获取数据计算结果
def calresult(h,hn,L1,L2,L3,L4):
    list = hn.tolist()
    list_new = []
    for i in list:
        if i != '':
            list_new.append(int(i))

    hn = np.array(list_new)
    try:
        max_deviation = abs(np.max(hn)-int(h)) #最大偏差
    except:
        max_deviation = False
    try:
        range = np.max(hn) - np.min(hn)
    except:
        range = False

    if L1!='' and L2!='':
        try:
            L1_2 = abs(int(L1) - int(L2))
        except:
            L1_2 = False
    else:
        L1_2 = False

    if L3!='' and L4!='':
        try:
            L3_4 = abs(int(L3) - int(L4))
        except:
            L3_4 = False
    else:
        L3_4 = False
    return max_deviation,range,L1_2,L3_4


#获取docx中表格数据
def get_data_docx(path,table_num):

    doc = Document(path)
    t = doc.tables[table_num]

    style = doc.styles['Normal']
    font = style.font
    font.name = u'新宋体'
    font.size = Pt(12)

    calculate_row_values =[]

    for i in range(4,12):
        index = 0
        data = []
        for j in range(1,13):
            cellv = t.cell(i, j).text

            if index in [5,9]:
                pass
            else:
                data.append(str(cellv))
            index += 1
        h = data[0]
        hn = np.array(data[1:6])
        L1 = data[6]
        L2 = data[7]
        L3 = data[8]
        L4 = data[9]
        calculate_row_values.append(calresult(h, hn, L1, L2, L3, L4))

    return calculate_row_values

#将计算结果写入docx文档中
def write_to_docx(path,textshow):
    doc = Document(path)
    style = doc.styles['Normal']
    font = style.font
    font.name = u'新宋体'
    font.size = Pt(12)
    for tablenum in range(len(doc.tables)):
        t = doc.tables[tablenum]
        get_calculatevalue = get_data_docx(path, tablenum)
        index = 0
        for r in range(4,12):
            mxdiff = get_calculatevalue[index][0]
            diff = get_calculatevalue[index][1]
            diff1 = get_calculatevalue[index][2]
            diff2 = get_calculatevalue[index][3]

            if mxdiff or str(mxdiff)=='0':
                if int(mxdiff)>99:
                    logging.info('页码:'+str(tablenum+1)+',净高最大偏差异常:'+str(mxdiff))
                    textshow.insert(END, '页码:'+str(tablenum+1)+',净高最大偏差异常:'+str(mxdiff))
                t.cell(r,13).text = str(mxdiff)
            else:
                t.cell(r,13).text = str('')

            if diff or str(diff)=='0':
                if int(diff) > 99:
                    logging.info('页码:'+str(tablenum + 1)+ ',净高极差异常:' + str(diff))
                    textshow.insert(END, '页码:'+str(tablenum + 1)+ ',净高极差异常:' + str(diff))

                t.cell(r,14).text = str(diff)
            else:
                t.cell(r,14).text = str('')

            if (diff1 or str(diff1)=='0') and (diff2 or str(diff2)=='0'):
                if int(diff1)>99:
                    logging.info('页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff1))
                    textshow.insert(END, '页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff1))


                if int(diff2) > 99:

                    logging.info('页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff2))
                    textshow.insert(END, '页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff2))


                t.cell(r,15).text = str(diff1)+'/'+str(diff2)
            elif diff1 or str(diff1)=='0':
                if int(diff1) > 99:
                    logging.info('页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff1))
                    textshow.insert(END, '页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff1))

                t.cell(r,15).text = str(diff1)+'/'
            elif diff2 or str(diff2)=='0':
                if int(diff2) > 99:
                    logging.info('页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff2))
                    textshow.insert(END, '页码:'+str(tablenum + 1)+ ',净开间差异常:' + str(diff2))

                t.cell(r,15).text = '/'+str(diff2)
            else:
                t.cell(r, 15).text = str('')
            index += 1
            if index >=8:
                break
    doc.save(path)


#GUI
class mainWindow(object):
    def __init__(self,path):
        self.path = path
        self.root = Tk()
        self.root.title(u'计算差值')
        # self.root.geometry("400x300")
        self.setCenter(400, 300)

        self.root.resizable(width=False, height=False)
        self.button = Button(self.root,text="开始",command=self.runCal)
        self.button.pack()

        self.button1 = Button(self.root,text='清空',command=self.clear).pack()
        self.label = Label(self.root,text='点击开始之前请关闭Word文档！').pack()
        self.text = Text(self.root)
        self.text.pack(side=LEFT,fill=BOTH,expand=1)
        self.text.focus_force()

        self.scroolbar = Scrollbar(self.root)
        self.text.config(yscrollcommand=self.scroolbar.set,width=20,height=20,background='#ffffff')

        self.scroolbar.config(command=self.text.yview)
        self.scroolbar.pack(side=RIGHT, fill=Y)

        self.root.mainloop()

    def clear(self):
        self.text.delete(1.0,END)

    def setCenter(self, width, height):

        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight() - 100

        self.root.update_idletasks()
        self.root.deiconify()
        self.root.withdraw()
        self.root.geometry('%sx%s+%s+%s' % (
            width + 10, height + 10, (screen_width - width) / 2,
            (screen_height - height) / 2))
        self.root.deiconify()

    #运行
    def runCal(self):
        if self.button['text'] == u'开始':
            self.button['text'] = '正在运行'
        else:
            self.text.insert(END, '正在运行，请等待！\n')
            return
        self.text.insert(END,'\n开始时间:'+start_time+'\n')
        try:
            t = ThreadDoc(write_to_docx,self.path,self.text,self.button)
            t.setDaemon(True)
            t.start()
        except:
            info = sys.exc_info()
            self.text.insert(END, str(info[0]) + ':' + str(info[1]) + '\n')
            self.text.insert(END, '无法正常运行(请先查看表格是否关闭)！\n')



if __name__ == "__main__":

    #针对打包文件,处理路径
    filename = (__file__).split('/')[-1].split('.')[0]
    if (filename + '.app') in os.getcwd().split(os.sep):
        pathlist = os.getcwd().split('.app')[0].split('/')[:-1]
        filepath = os.sep + os.path.join(*pathlist) + os.sep + 'source.docx'
    else:
        filepath = os.getcwd() + os.sep + 'source.docx'
    logging.info(filepath)
    mw = mainWindow(filepath)